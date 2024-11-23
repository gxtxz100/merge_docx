from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import sys

def copy_paragraph_format(new_paragraph, paragraph):
    """复制段落格式"""
    # 复制段落级别的格式
    if hasattr(paragraph, 'style') and paragraph.style:
        try:
            new_paragraph.style = paragraph.style
        except:
            pass
            
    # 复制段落格式属性
    pf = new_paragraph.paragraph_format
    old_pf = paragraph.paragraph_format
    
    # 复制具体的段落格式
    attributes = [
        'alignment', 'line_spacing', 'space_before', 'space_after',
        'left_indent', 'right_indent', 'first_line_indent'
    ]
    
    for attr in attributes:
        try:
            if hasattr(old_pf, attr) and getattr(old_pf, attr) is not None:
                setattr(pf, attr, getattr(old_pf, attr))
        except:
            continue

def copy_run_format(new_run, run):
    """复制运行格式，确保字体格式完整复制"""
    # 复制基本属性
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    
    # 复制字体属性
    if run._element.rPr is not None:
        # 如果新run没有rPr元素，创建一个
        if new_run._element.rPr is None:
            new_run._element.get_or_add_rPr()
            
        # 复制字体名称（包括中文字体和西文字体）
        if run.font.name:
            # 设置中文字体
            r = new_run._element.rPr
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:eastAsia'), run.font.name)  # 中文字体
            fonts.set(qn('w:ascii'), run.font.name)     # 英文字体
            fonts.set(qn('w:hAnsi'), run.font.name)     # 其他文字
            r.append(fonts)
            
        # 复制字体大小
        if run.font.size:
            new_run.font.size = run.font.size
            
        # 复制字体颜色
        if run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
            
        # 复制删除线
        try:
            new_run.font.strike = run.font.strike
        except:
            pass

def natural_sort_key(s):
    """实现 Windows 风格的自然排序"""
    import re
    # 将字符串分割成文本和数字部分
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split('([0-9]+)', str(s))]

def merge_docx_files(folder_path):
    # 创建一个新的文档作为合并后的文档
    merged_doc = Document()
    
    # 获取文件夹中所有的.docx文件并按 Windows 方式排序
    folder = Path(folder_path)
    docx_files = sorted(folder.glob('*.docx'), key=lambda x: natural_sort_key(x.name))
    
    if not docx_files:
        print("错误：在指定文件夹中没有找到.docx文件")
        return
    
    # 首先复制第一个文档的样式
    first_doc = Document(docx_files[0])
    # 复制样式
    for style in first_doc.styles:
        if style.name not in merged_doc.styles:
            try:
                merged_doc.styles.add_style(style.name, style.type, True)
            except:
                pass
    
    # 遍历所有docx文件并合并
    for i, file_path in enumerate(docx_files):
        # 跳过输出文件（如果存在）
        if file_path.name == 'output.docx':
            continue
            
        print(f"正在处理文件: {file_path.name}")
        
        # 打开当前文档
        doc = Document(file_path)
        
        # 如果不是第一个文档，添加分页符
        if i > 0:
            merged_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        # 复制所有段落
        for paragraph in doc.paragraphs:
            new_paragraph = merged_doc.add_paragraph()
            copy_paragraph_format(new_paragraph, paragraph)
            
            # 复制所有运行及其格式
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                copy_run_format(new_run, run)
        
        # 复制所有表格
        for table in doc.tables:
            # 创建新表格
            new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            # 复制表格样式
            if table.style:
                new_table.style = table.style
            
            # 复制单元格内容和格式
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.cell(i, j)
                    # 复制单元格中的段落
                    for paragraph in cell.paragraphs:
                        new_paragraph = new_cell.add_paragraph()
                        copy_paragraph_format(new_paragraph, paragraph)
                        for run in paragraph.runs:
                            new_run = new_paragraph.add_run(run.text)
                            copy_run_format(new_run, run)

    # 保存合并后的文档
    output_path = folder / 'output.docx'
    merged_doc.save(output_path)
    print(f"\n合并完成！输出文件保存在: {output_path}")

def main():
    # 获取用户输入的文件夹路径
    if len(sys.argv) > 1:
        folder_path = sys.argv[1]
    else:
        folder_path = input("请输入包含Word文档的文���夹路径: ")
    
    # 检查文件夹是否存在
    if not Path(folder_path).exists():
        print("错误：指定的文件夹不存在")
        return
    
    try:
        merge_docx_files(folder_path)
    except Exception as e:
        print(f"发生错误: {str(e)}")

if __name__ == '__main__':
    main()
